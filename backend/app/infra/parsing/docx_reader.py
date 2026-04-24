from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document as DocxDocument

from .base import BaseParser
from .pymupdf_reader import ParsedDocument


def _table_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm: list[list[str]] = []
    for r in rows:
        rr = [str(c).strip() for c in r]
        while len(rr) < width:
            rr.append("")
        norm.append(rr)
    header = norm[0]
    body = norm[1:]
    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


class DocxReader(BaseParser):
    def parse(self, data: bytes, params: dict[str, Any]) -> ParsedDocument:
        # python-docx supports .docx; legacy .doc should be converted before ingest.
        try:
            doc = DocxDocument(BytesIO(data))
        except Exception as e:
            raise RuntimeError("Failed to parse Word file. Only .docx is supported (please convert .doc → .docx).") from e
        parts: list[str] = []

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([(cell.text or "").strip() for cell in row.cells])
            md = _table_to_markdown(rows)
            if md.strip():
                parts.append(md)

        text = "\n\n".join(parts).strip()
        if not text:
            raise RuntimeError("Parsed document is empty")
        return ParsedDocument(text=text, page_count=1)

